import os
import re
import sqlite3
import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
from datetime import datetime
from telebot.types import Message
from docx import Document
import openpyxl
import phonenumbers
from phonenumbers import geocoder, carrier
import time
import csv
import sys
from pathlib import Path
from tqdm import tqdm
import io
import threading
from memepay import AsyncMemePay, MemePay
from memepay.webhook import WebhookPayload
import asyncio
from telebot import util
from telebot.apihelper import ApiTelegramException
import hmac
import hashlib
import requests
from urllib.parse import urlencode
import uuid
import schedule 


TOKEN = '8348293204:AAHZ49kQeXwARjEfMOuIv5-lFpjeNhshNqg'  # Замените на свой токен
bot = telebot.TeleBot(TOKEN)

BASE_DIR = 'Base'

ADMINS = [7209314948, 1811753312, 8441679190]

MEMEPAY_API_KEY = "mp_9992c136f908ac7d2467c3b197d83cfc"  # Получить в личном кабинете MemePay
MEMEPAY_SHOP_ID = "cb457a74-17ca-4365-91f6-51f0886679c4"   # ID магазина в MemePay

LOG_GROUP_ID = -1002991202021  # ID группы для логов





# --- Импорт данных из файла ---
def extract_data_from_csv_file(file_path):
    names, dob, phones, emails, addresses, tg_ids, tg_usernames = [], [], [], [], [], [], []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f, delimiter=';')
            first_line = next(reader)

            # Проверяем, есть ли Telegram данные в заголовках
            if 'id' in first_line and 'username' in first_line:
                # Обработка Telegram формата (id,phone,username,first_name,last_name)
                headers = [h.strip().lower() for h in first_line]
                header_map = {h: i for i, h in enumerate(headers)}

                idx_id = header_map.get('id', None)
                idx_phone = header_map.get('phone', None)
                idx_username = header_map.get('username', None)
                idx_first_name = header_map.get('first_name', None)
                idx_last_name = header_map.get('last_name', None)

                for row in reader:
                    if len(row) < len(headers):
                        continue
                    
                    # Telegram данные
                    if idx_id is not None and idx_id < len(row):
                        tg_ids.append(row[idx_id].strip())
                    if idx_username is not None and idx_username < len(row):
                        tg_usernames.append(row[idx_username].strip())
                    
                    # Телефон и имя
                    if idx_phone is not None and idx_phone < len(row):
                        phones.append(row[idx_phone].strip())
                    if idx_first_name is not None and idx_first_name < len(row):
                        first_name = row[idx_first_name].strip()
                        last_name = row[idx_last_name].strip() if idx_last_name is not None and idx_last_name < len(row) else ''
                        names.append(f"{first_name} {last_name}".strip())

            # Проверяем, стандартные ли заголовки в первой строке
            elif any(h.lower() in ['фамилия', 'имя', 'отчество', 'дата рождения', 'телефон'] for h in first_line):
                headers = first_line
                header_map = {h.strip(): i for i, h in enumerate(headers)}

                idx_lastname = header_map.get('Фамилия', None)
                idx_firstname = header_map.get('Имя', None)
                idx_middlename = header_map.get('Отчество', None)
                idx_dob = header_map.get('Дата рождения', None)
                idx_phone = header_map.get('Телефон', None)
                idx_email = header_map.get('Электронная почта', None)
                idx_address = header_map.get('Адрес', None)

                for row in reader:
                    if len(row) < len(headers):
                        continue
                    lastname = row[idx_lastname].strip() if idx_lastname is not None and idx_lastname < len(row) else ''
                    firstname = row[idx_firstname].strip() if idx_firstname is not None and idx_firstname < len(row) else ''
                    middlename = row[idx_middlename].strip() if idx_middlename is not None and idx_middlename < len(row) else ''
                    full_name = ' '.join([lastname, firstname, middlename]).strip()
                    names.append(full_name if full_name else '')
                    dob.append(row[idx_dob].strip() if idx_dob is not None and idx_dob < len(row) else '')
                    phones.append(row[idx_phone].strip() if idx_phone is not None and idx_phone < len(row) else '')
                    emails.append(row[idx_email].strip() if idx_email is not None and idx_email < len(row) else '')
                    addresses.append(row[idx_address].strip() if idx_address is not None and idx_address < len(row) else '')

            else:
                # Обработка строк без заголовков
                def process_row(row):
                    phone = row[0].strip() if len(row) > 0 else ''
                    name = row[1].strip() if len(row) > 1 else ''
                    birth = row[2].strip() if len(row) > 2 else ''
                    address = ', '.join(x.strip() for x in row[4:] if x.strip()) if len(row) > 4 else ''
                    return name, birth, phone, '', address

                name, birth, phone, email, address = process_row(first_line)
                names.append(name)
                dob.append(birth)
                phones.append(phone)
                emails.append(email)
                addresses.append(address)

                for row in reader:
                    name, birth, phone, email, address = process_row(row)
                    names.append(name)
                    dob.append(birth)
                    phones.append(phone)
                    emails.append(email)
                    addresses.append(address)

        return names, dob, phones, emails, addresses, tg_ids, tg_usernames

    except Exception as e:
        print(f"❌ Ошибка при обработке CSV: {e}")
        return [], [], [], [], [], [], []


def is_admin(user_id):
    ADMINS = [7209314948, 1811753312, 8441679190]  # Замените на ваш Telegram ID
    return user_id in ADMINS


def extract_data_from_text(text):
    # Проверка на CSV с заголовками
    if 'Дата рождения' in text and 'Электронная почта' in text:
        f = io.StringIO(text.replace('"', ''))  # убрать кавычки
        reader = csv.reader(f, delimiter=';')
        headers = next(reader)

        header_map = {h.strip(): i for i, h in enumerate(headers)}

        idx_lastname = header_map.get('Фамилия', None)
        idx_firstname = header_map.get('Имя', None)
        idx_middlename = header_map.get('Отчество', None)
        idx_dob = header_map.get('Дата рождения', None)
        idx_phone = header_map.get('Телефон', None)
        idx_email = header_map.get('Электронная почта', None)
        idx_address = header_map.get('Адрес', None)

        names, dob, phones, emails, addresses = [], [], [], [], []

        for row in reader:
            if len(row) < len(headers):
                continue  # строка неполная

            lastname = row[idx_lastname].strip() if idx_lastname is not None and idx_lastname < len(row) else ''
            firstname = row[idx_firstname].strip() if idx_firstname is not None and idx_firstname < len(row) else ''
            middlename = row[idx_middlename].strip() if idx_middlename is not None and idx_middlename < len(row) else ''
            full_name = ' '.join([lastname, firstname, middlename]).strip()
            names.append(full_name if full_name else '')

            dob.append(row[idx_dob].strip() if idx_dob is not None and idx_dob < len(row) else '')
            phones.append(row[idx_phone].strip() if idx_phone is not None and idx_phone < len(row) else '')
            emails.append(row[idx_email].strip() if idx_email is not None and idx_email < len(row) else '')
            addresses.append(row[idx_address].strip() if idx_address is not None and idx_address < len(row) else '')

        return names, dob, phones, emails, addresses

    # иначе обычный парсинг
    phones = re.findall(r'\+?\d{10,15}', text)
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)
    names = re.findall(r'[А-ЯЁA-Z][а-яёa-z]+ [А-ЯЁA-Z][а-яёa-z]+(?: [А-ЯЁA-Z][а-яёa-z]+)?', text)
    dob = re.findall(r'\d{2}[./-]\d{2}[./-]\d{4}', text)
    addresses = []
    return names, dob, phones, emails, addresses


def colorize_bar(bar, percent):
    if percent < 30:
        color = '\033[91m'  # 🔴 Красный
    elif percent < 70:
        color = '\033[93m'  # 🟡 Жёлтый
    else:
        color = '\033[92m'  # 🟢 Зелёный
    reset = '\033[0m'
    return f"{color}{bar}{reset}"


def parse_file(file_path):
    try:
        content = ''
        file_size = os.path.getsize(file_path)
        read_size = 0

        def print_progress(percent):
            bar_length = 30
            filled_length = int(bar_length * percent // 100)
            bar = '█' * filled_length + '░' * (bar_length - filled_length)
            sys.stdout.write(f"\r📄 Чтение: {colorize_bar(f'[{bar}]', percent)} {percent}%")
            sys.stdout.flush()

        if file_path.endswith('.txt') or file_path.endswith('.csv'):
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    content += line
                    read_size += len(line.encode('utf-8'))
                    percent = int((read_size / file_size) * 100)
                    print_progress(percent)

        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            total = len(doc.paragraphs)
            for i, p in enumerate(doc.paragraphs):
                content += p.text + '\n'
                percent = int((i + 1) / total * 100)
                print_progress(percent)

        elif file_path.endswith('.xlsx'):
            wb = openpyxl.load_workbook(file_path, read_only=True)
            total = sum(sheet.max_row for sheet in wb.worksheets)
            count = 0
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    line = ' '.join([str(cell) for cell in row if cell]) + '\n'
                    content += line
                    count += 1
                    percent = int((count / total) * 100)
                    print_progress(percent)
        else:
            print(f"⚠️ Формат файла {file_path} не поддерживается.")

        print("\n✅ Файл прочитан.")
        return content

    except Exception as e:
        print(f"\n❌ Ошибка при чтении файла {file_path}: {e}")
        return ''


def create_database(folder):
    os.makedirs(folder, exist_ok=True)
    conn = sqlite3.connect(os.path.join(folder, 'base.db'))
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS people (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        dob TEXT,
        phone TEXT,
        email TEXT,
        address TEXT,
        tg_id TEXT,
        tg_username TEXT
    )''')
    conn.commit()
    return conn


def import_data_to_db(conn, names, dob, phones, emails, addresses, tg_ids, tg_usernames):
    try:
        c = conn.cursor()
        max_len = max(len(names), len(dob), len(phones), len(emails), len(tg_ids), len(tg_usernames))
        if max_len == 0:
            print("❌ Нет данных для импорта.")
            return

        start_time = time.time()

        for i in range(max_len):
            c.execute("INSERT INTO people (name, dob, phone, email, address, tg_id, tg_username) VALUES (?, ?, ?, ?, ?, ?, ?)", (
                names[i] if i < len(names) else '',
                dob[i] if i < len(dob) else '',
                phones[i] if i < len(phones) else '',
                emails[i] if i < len(emails) else '',
                addresses[i] if i < len(addresses) else '',
                tg_ids[i] if i < len(tg_ids) else '',
                tg_usernames[i] if i < len(tg_usernames) else ''
            ))

            if (i + 1) % 100 == 0 or i == max_len - 1:
                conn.commit()

            # Прогресс
            percent = int((i + 1) / max_len * 100)
            elapsed = time.time() - start_time
            speed = (i + 1) / elapsed if elapsed > 0 else 0
            sys.stdout.write(
                f"\r📥 Импорт: {i + 1}/{max_len} записей | {percent}% | Скорость: {speed:.2f} зап/сек"
            )
            sys.stdout.flush()

        print("\n✅ Импорт завершён успешно.")
    except Exception as e:
        print(f"\n❌ Ошибка при импорте в базу: {e}")


def confirm_import():
    choice = input("Импортировать найденные данные? (да/нет): ").strip().lower()
    return choice == 'да'


def import_database():
    print("📁 Доступные файлы для импорта в текущей папке:")
    files = [f for f in os.listdir() if f.endswith(('.txt', '.docx', '.xlsx', '.csv'))]
    if not files:
        print("Файлов для импорта нет.")
        return
    
    for i, f in enumerate(files, 1):
        print(f"{i}. {f}")

    # Выбор файла
    try:
        file_num = int(input("Введите номер файла для импорта: ").strip())
        if file_num < 1 or file_num > len(files):
            print("❌ Неверный номер файла.")
            return
        filename = files[file_num - 1]
    except ValueError:
        print("❌ Введите число.")
        return

    # Выбор типа базы
    print("\nВыберите тип базы:")
    print("1. Обычная база (ФИО, телефон, дата рождения и т.д.)")
    print("2. Telegram база (ID, phone, username, first_name, last_name)")
    print("3. Почты и пароли (email:password)")
    print("4. Почта;ФИО;ДР;Работодатель;Адрес")
    print("5. ID|Full name|Birthday|Phone| (формат с разделителем |)")  # НОВЫЙ ТИП
    
    try:
        base_type = int(input("Введите номер типа (1, 2, 3, 4 или 5): ").strip())
        if base_type not in [1, 2, 3, 4, 5]:  # Добавили 5
            print("❌ Неверный тип базы.")
            return
    except ValueError:
        print("❌ Введите число.")
        return

    if filename.endswith('.csv'):
        if base_type == 1:
            names, dob, phones, emails, addresses = extract_regular_csv(filename)
            tg_ids, tg_usernames = [], []
        elif base_type == 2:
            names, dob, phones, emails, addresses, tg_ids, tg_usernames = extract_telegram_csv(filename)
        elif base_type == 3:
            emails, passwords = extract_email_password_csv(filename)
            names, dob, phones, addresses, tg_ids, tg_usernames = [], [], [], [], [], []
        elif base_type == 4:
            emails, names, dob, addresses = extract_email_fio_dob_address_csv(filename)
            phones, tg_ids, tg_usernames = [], [], []
        elif base_type == 5:  # НОВЫЙ ТИП
            names, dob, phones = extract_pipe_format_csv(filename)
            emails, addresses, tg_ids, tg_usernames = [], [], [], []
    else:
        # Для других форматов файлов
        content = parse_file(filename)
        if base_type == 1:
            names, dob, phones, emails, addresses = extract_data_from_text(content)
            tg_ids, tg_usernames = [], []
        elif base_type == 2:
            print("❌ Для этого типа файлов поддерживается только обычная база.")
            return
        elif base_type == 3:
            emails, passwords = extract_email_password_text(content)
            names, dob, phones, addresses, tg_ids, tg_usernames = [], [], [], [], [], []
        elif base_type == 4:
            emails, names, dob, addresses = extract_email_fio_dob_address_text(content)
            phones, tg_ids, tg_usernames = [], [], []
        elif base_type == 5:  # НОВЫЙ ТИП
            names, dob, phones = extract_pipe_format_text(content)
            emails, addresses, tg_ids, tg_usernames = [], [], [], []

    # Остальной код функции остается без изменений...
    print("\n📊 Найдено:")
    if base_type == 3:
        print(f"├📧 Почты: {len(emails)}")
        print(f"├🔑 Пароли: {len(passwords)}")
        print(f"╰📦 Всего записей: {len(emails)}")
    elif base_type == 4:
        print(f"├📧 Почты: {len(emails)}")
        print(f"├👤 ФИО: {len(names)}")
        print(f"├🎂 Дата рождения: {len(dob)}")
        print(f"├📍 Адреса: {len(addresses)}")
        print(f"╰📦 Всего записей: {max(len(emails), len(names), len(dob), len(addresses))}")
    elif base_type == 5:  # НОВЫЙ ТИП
        print(f"├👤 ФИО: {len(names)}")
        print(f"├🎂 Дата рождения: {len(dob)}")
        print(f"├📞 Телефоны: {len(phones)}")
        print(f"╰📦 Всего записей: {max(len(names), len(dob), len(phones))}")
    else:
        print(f"├👤 ФИО: {len(names)}")
        print(f"├🎂 Дата рождения: {len(dob)}")
        print(f"├📞 Телефоны: {len(phones)}")
        print(f"├📧 Почты: {len(emails)}")
        if base_type == 2:
            print(f"├🆔 Telegram ID: {len(tg_ids)}")
            print(f"├👤 Telegram юзернеймы: {len(tg_usernames)}")
        print(f"╰📦 Всего записей: {max(len(names), len(dob), len(phones), len(emails), len(tg_ids), len(tg_usernames))}")

    if confirm_import():
        base_name = os.path.splitext(os.path.basename(filename))[0]
        base_path = os.path.join(BASE_DIR, base_name)
        conn = create_database(base_path)
        
        if base_type == 3:
            import_data_to_db(conn, names, dob, phones, emails, addresses, tg_ids, tg_usernames)
            create_email_password_table(conn)
            import_email_password_data(conn, emails, passwords)
        elif base_type == 4:
            import_data_to_db(conn, names, dob, phones, emails, addresses, tg_ids, tg_usernames)
        elif base_type == 5:  # НОВЫЙ ТИП
            import_data_to_db(conn, names, dob, phones, emails, addresses, tg_ids, tg_usernames)
        else:
            import_data_to_db(conn, names, dob, phones, emails, addresses, tg_ids, tg_usernames)
            
        conn.close()
        print("✅ Импорт завершён!")
    else:
        print("❌ Импорт отменён.")

def extract_telegram_csv(file_path):
    """Специальная функция для извлечения данных из Telegram CSV в формате 'id|name|fname|phone|uid|nik|wo'"""
    names, dob, phones, emails, addresses, tg_ids, tg_usernames = [], [], [], [], [], [], []
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                parts = line.strip().split('|')
                if len(parts) < 6:  # Минимально нужные поля: phone и uid
                    continue
                
                # Обработка номера телефона
                phone = parts[3].strip() if len(parts) > 3 else ''
                if phone:
                    # Нормализация российских номеров (добавляем +7 если номер начинается с 7 или 8)
                    if phone.startswith('7') and len(phone) == 11:
                        phone = f"+{phone}"
                    elif phone.startswith('8') and len(phone) == 11:
                        phone = f"+7{phone[1:]}"
                    # Удаляем все нецифровые символы
                    phone = ''.join(c for c in phone if c.isdigit() or c == '+')
                
                # Telegram ID
                tg_id = parts[4].strip() if len(parts) > 4 else ''
                
                # Обработка username (никнейма)
                username = parts[5].strip() if len(parts) > 5 else ''
                if username and not username.startswith('@'):
                    username = f"@{username}"
                
                # Формируем имя из name и fname
                name_parts = []
                if len(parts) > 1 and parts[1].strip():
                    name_parts.append(parts[1].strip())
                if len(parts) > 2 and parts[2].strip():
                    name_parts.append(parts[2].strip())
                name = ' '.join(name_parts) if name_parts else ''
                
                # Добавляем данные в списки
                if phone:
                    phones.append(phone)
                if tg_id:
                    tg_ids.append(tg_id)
                if username:
                    tg_usernames.append(username)
                if name:
                    names.append(name)
                
                # Пустые поля для Telegram базы
                dob.append('')
                emails.append('')
                addresses.append('')
                
        return names, dob, phones, emails, addresses, tg_ids, tg_usernames
    except Exception as e:
        print(f"❌ Ошибка при обработке Telegram CSV: {e}")
        return [], [], [], [], [], [], []

def extract_regular_csv(file_path):
    """Функция для извлечения данных из обычного CSV"""
    # Здесь оставляем вашу оригинальную функцию extract_data_from_csv_file,
    # но возвращаем только 5 значений и два пустых списка для совместимости
    names, dob, phones, emails, addresses = extract_data_from_csv_file(file_path)
    return names, dob, phones, emails, addresses, [], []

def extract_email_password_csv(file_path):
    """Извлекает почты и пароли из CSV файла в формате email:password"""
    emails = []
    passwords = []
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) >= 1:
                    line = row[0].strip()
                    if ':' in line:
                        email, password = line.split(':', 1)
                        emails.append(email.strip())
                        passwords.append(password.strip())
                    elif '@' in line and any(c in line for c in ['|', ';', ',']):
                        # Пробуем другие разделители
                        for sep in ['|', ';', ',']:
                            if sep in line:
                                parts = line.split(sep, 1)
                                if len(parts) == 2 and '@' in parts[0]:
                                    emails.append(parts[0].strip())
                                    passwords.append(parts[1].strip())
                                    break
        return emails, passwords
    except Exception as e:
        print(f"❌ Ошибка при обработке CSV с почтами: {e}")
        return [], []

def extract_email_password_text(text):
    """Извлекает почты и пароли из текста"""
    emails = []
    passwords = []
    
    # Ищем строки в формате email:password
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if ':' in line and '@' in line.split(':')[0]:
            email, password = line.split(':', 1)
            emails.append(email.strip())
            passwords.append(password.strip())
        elif '|' in line and '@' in line.split('|')[0]:
            email, password = line.split('|', 1)
            emails.append(email.strip())
            passwords.append(password.strip())
        elif ';' in line and '@' in line.split(';')[0]:
            email, password = line.split(';', 1)
            emails.append(email.strip())
            passwords.append(password.strip())
        elif re.match(r'[\w\.-]+@[\w\.-]+', line) and len(line.split()) >= 2:
            # Если почта и пароль разделены пробелом
            parts = line.split()
            for i, part in enumerate(parts):
                if '@' in part:
                    emails.append(part.strip())
                    if i + 1 < len(parts):
                        passwords.append(parts[i + 1].strip())
                    else:
                        passwords.append('')
                    break
    
    return emails, passwords

def create_email_password_table(conn):
    """Создает таблицу для хранения почт и паролей"""
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS emails (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT UNIQUE,
        password TEXT
    )''')
    conn.commit()

def import_email_password_data(conn, emails, passwords):
    """Импортирует данные почт и паролей в базу"""
    try:
        c = conn.cursor()
        max_len = max(len(emails), len(passwords))
        
        for i in range(max_len):
            email = emails[i] if i < len(emails) else ''
            password = passwords[i] if i < len(passwords) else ''
            
            if email:  # Импортируем только если есть почта
                c.execute("INSERT OR IGNORE INTO emails (email, password) VALUES (?, ?)", 
                         (email, password))
            
            if (i + 1) % 100 == 0:
                conn.commit()
                print(f"📧 Импортировано {i + 1}/{max_len} записей")
        
        conn.commit()
        print(f"✅ Импорт почт завершен: {len(emails)} записей")
        
    except Exception as e:
        print(f"❌ Ошибка при импорте почт: {e}")

def extract_email_fio_dob_address_csv(file_path):
    """Извлекает данные из CSV в формате: ПОЧТА;ФИО;ДР;Работодатель;Адрес"""
    emails, names, dob, addresses = [], [], [], []
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f, delimiter=';')
            
            # Пропускаем заголовок если есть
            first_line = next(reader, None)
            if first_line and 'ПОЧТА' in first_line[0] and 'ФИО' in first_line[1]:
                print("📋 Обнаружен заголовок, пропускаем...")
            else:
                # Если нет заголовка, возвращаемся к первой строке
                f.seek(0)
                reader = csv.reader(f, delimiter=';')
            
            for row in reader:
                if len(row) < 5:  # Минимум 5 полей
                    continue
                
                # Извлекаем данные (игнорируем работодателя)
                email = row[0].strip() if row[0].strip() else ''
                fio = row[1].strip() if row[1].strip() else ''
                birth_date = row[2].strip() if row[2].strip() else ''
                address = row[4].strip() if len(row) > 4 and row[4].strip() else ''
                
                if email:
                    emails.append(email)
                    names.append(fio)
                    dob.append(birth_date)
                    addresses.append(address)
                
        return emails, names, dob, addresses
        
    except Exception as e:
        print(f"❌ Ошибка при обработке CSV формата 'ПОЧТА;ФИО;ДР;Работодатель;Адрес': {e}")
        return [], [], [], []
    
def extract_email_fio_dob_address_text(text):
    """Извлекает данные из текста в формате: ПОЧТА;ФИО;ДР;Работодатель;Адрес"""
    emails, names, dob, addresses = [], [], [], []
    
    lines = text.split('\n')
    for line in lines:
        parts = line.split(';')
        if len(parts) >= 5:
            email = parts[0].strip()
            fio = parts[1].strip()
            birth_date = parts[2].strip()
            address = parts[4].strip() if len(parts) > 4 else ''
            
            if email and '@' in email:
                emails.append(email)
                names.append(fio)
                dob.append(birth_date)
                addresses.append(address)
    
    return emails, names, dob, addresses

def extract_pipe_format_csv(file_path):
    """Извлекает данные из CSV в формате: ID|Full name|Birthday|Phone|"""
    names, dob, phones = [], [], []
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                parts = line.strip().split('|')
                if len(parts) < 4:  # Минимум 4 поля
                    continue
                
                # Извлекаем данные
                full_name = parts[1].strip() if len(parts) > 1 else ''
                birth_date = parts[2].strip() if len(parts) > 2 else ''
                phone = parts[3].strip() if len(parts) > 3 else ''
                
                # Обрабатываем поле Full name - извлекаем ФИО из скобок если есть
                if '(' in full_name and ')' in full_name:
                    # Ищем ФИО в скобках
                    match = re.search(r'\((.*?)\)', full_name)
                    if match:
                        full_name = match.group(1).strip()
                
                # Нормализуем телефон
                if phone and phone.lower() != 'null':
                    phone = normalize_phone(phone)
                else:
                    phone = ''
                
                if full_name or birth_date or phone:
                    names.append(full_name)
                    dob.append(birth_date)
                    phones.append(phone)
                
        return names, dob, phones
        
    except Exception as e:
        print(f"❌ Ошибка при обработке CSV формата 'ID|Full name|Birthday|Phone|': {e}")
        return [], [], []

def extract_pipe_format_text(text):
    """Извлекает данные из текста в формате: ID|Full name|Birthday|Phone|"""
    names, dob, phones = [], [], []
    
    lines = text.split('\n')
    for line in lines:
        parts = line.split('|')
        if len(parts) >= 4:
            full_name = parts[1].strip() if len(parts) > 1 else ''
            birth_date = parts[2].strip() if len(parts) > 2 else ''
            phone = parts[3].strip() if len(parts) > 3 else ''
            
            # Обрабатываем поле Full name - извлекаем ФИО из скобок если есть
            if '(' in full_name and ')' in full_name:
                match = re.search(r'\((.*?)\)', full_name)
                if match:
                    full_name = match.group(1).strip()
            
            # Нормализуем телефон
            if phone and phone.lower() != 'null':
                phone = normalize_phone(phone)
            else:
                phone = ''
            
            if full_name or birth_date or phone:
                names.append(full_name)
                dob.append(birth_date)
                phones.append(phone)
    
    return names, dob, phones


# --- Поиск и форматирование ---

def normalize_phone(phone):
    if not phone:
        return ""
    
    # Убираем все нецифровые символы кроме плюса
    phone = re.sub(r'[^\d+]', '', str(phone))
    
    # Обработка российских номеров
    if phone.startswith('8') and len(phone) == 11:
        return '+7' + phone[1:]
    elif phone.startswith('7') and len(phone) == 11:
        return '+' + phone
    elif len(phone) == 10 and not phone.startswith('+'):
        return '+7' + phone
    
    return phone


def search_number_variants(conn, phone):
    phone_plus = normalize_phone(phone)
    phone_plain = phone_plus.lstrip('+')

    c = conn.cursor()
    c.execute("SELECT * FROM people WHERE phone LIKE ? OR phone LIKE ?", (f'%{phone_plus}%', f'%{phone_plain}%'))
    return c.fetchall()


def get_phone_info(phone):
    try:
        parsed = phonenumbers.parse(normalize_phone(phone), None)
        country = geocoder.country_name_for_number(parsed, "ru")
        region = geocoder.description_for_number(parsed, "ru")
        operator = carrier.name_for_number(parsed, "ru")
        country = country if country else None
        region = region if region else None
        operator = operator if operator else None
        return operator, region, country
    except:
        return None, None, None


def calculate_age(dob):
    try:
        # Поддержка нескольких форматов даты
        for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d-%m-%Y"):
            try:
                dob_date = datetime.strptime(dob, fmt)
                break
            except ValueError:
                continue
        else:
            return None  # Не распознали дату

        today = datetime.today()
        age = today.year - dob_date.year - ((today.month, today.day) < (dob_date.month, dob_date.day))
        return age
    except:
        return None




def format_report(record):
    # --- Почтовая база (emails) ---
    if len(record) == 3 and '@' in str(record[1]):
        _, email, password = record
        lines = []
        lines.append("📧 Почтовые данные")
        lines.append(f"├ Email: {email}")
        if password:
            lines.append(f"╰ Пароль: {password}")
        return "\n".join(lines)

    # --- База people (с tg_id и tg_username) ---
    if len(record) == 8:
        _, name, dob, phone, email, address, tg_id, tg_username = record
    else:  # Старый формат (без tg_id, tg_username)
        _, name, dob, phone, email, address = record
        tg_id, tg_username = None, None

    operator, region, country = get_phone_info(phone) if phone else (None, None, None)
    age = calculate_age(dob) if dob else None

    lines = []

    if phone:
        lines.append("📱 ")
        lines.append(f"├ Телефон: {phone}")
        if operator:
            lines.append(f"├ Оператор: {operator}")
        if region:
            lines.append(f"├ Регион: {region}")
        if country:
            lines.append(f"╰ Страна: {country}")

    if name or dob:
        lines.append("\n👤 Основные данные")
        if name:
            lines.append(f"├ ФИО: {name}")
        if dob:
            lines.append(f"├ Дата рождения: {dob}")
            if age is not None:
                lines.append(f"╰ Возраст: {age}")

    if email:
        lines.append(f"\n╰📧 E-mail: {email}")

    if address:
        lines.append(f"\n╰📍 Примерный адрес: {address}")

    if tg_id or tg_username:
        lines.append("\n🔹 Telegram данные")
        if tg_id:
            lines.append(f"├ ID: {tg_id}")
        if tg_username:
            username = f"@{tg_username}" if tg_username and not tg_username.startswith('@') else tg_username
            lines.append(f"╰ Username: {username}")

    if not (operator or region or country or tg_id or tg_username):
        lines.append("\n❗ По этому запросу нет дополнительной информации.")

    return "\n".join(lines)


def search_all_databases_by_phone(phone):
    results = []
    
    # Нормализуем номер
    phone = phone.strip().replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
    
    print(f"🔍 Исходный номер для поиска: {phone}")
    
    # Создаем варианты номеров для поиска
    search_variants = set()
    
    # Добавляем оригинальный номер
    search_variants.add(phone)
    
    # Обработка российских номеров
    if phone.startswith('+7') and len(phone) == 12:
        # +79500467560 → варианты: +79500467560, 79500467560, 89500467560, 9500467560
        clean_phone = phone[2:]  # 9500467560
        search_variants.update([
            phone,                    # +79500467560
            '7' + clean_phone,       # 79500467560
            '8' + clean_phone,       # 89500467560
            clean_phone              # 9500467560
        ])
    elif phone.startswith('7') and len(phone) == 11:
        # 79500467560 → варианты: +79500467560, 79500467560, 89500467560, 9500467560
        search_variants.update([
            '+7' + phone[1:],        # +79500467560
            phone,                    # 79500467560
            '8' + phone[1:],         # 89500467560
            phone[1:]                # 9500467560
        ])
    elif phone.startswith('8') and len(phone) == 11:
        # 89500467560 → варианты: +79500467560, 79500467560, 89500467560, 9500467560
        search_variants.update([
            '+7' + phone[1:],        # +79500467560
            '7' + phone[1:],         # 79500467560
            phone,                    # 89500467560
            phone[1:]                # 9500467560
        ])
    elif len(phone) == 10 and not phone.startswith(('+', '7', '8')):
        # 9500467560 → варианты: +79500467560, 79500467560, 89500467560, 9500467560
        search_variants.update([
            '+7' + phone,            # +79500467560
            '7' + phone,             # 79500467560
            '8' + phone,             # 89500467560
            phone                    # 9500467560
        ])
    
    # Фильтруем пустые значения
    search_variants = [v for v in search_variants if v]
    
    print(f"🔍 Варианты для поиска: {search_variants}")
    
    for folder in os.listdir(BASE_DIR):
        # Пропускаем Telegram-папки для поиска по телефону
        if 'telegram' in folder.lower() or 'глаз' in folder.lower():
            continue
            
        db_path = os.path.join(BASE_DIR, folder, 'base.db')
        if not os.path.exists(db_path):
            continue
            
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Ищем по всем вариантам
        for variant in search_variants:
            # Ищем точное совпадение
            c.execute("SELECT * FROM people WHERE phone = ?", (variant,))
            exact_matches = c.fetchall()
            
            # Ищем частичное совпадение (если номер хранится в составе строки)
            c.execute("SELECT * FROM people WHERE phone LIKE ?", (f'%{variant}%',))
            partial_matches = c.fetchall()
            
            # Объединяем результаты
            all_matches = exact_matches + partial_matches
            
            for record in all_matches:
                if (folder, record) not in results:  # избегаем дубликатов
                    results.append((folder, record))
        
        conn.close()
    
    print(f"🔍 Найдено результатов: {len(results)}")
    return results



def search_all_databases_by_query(query):
    query = query.strip()
    results = []
    
    print(f"🔍 Общий поиск по запросу: '{query}'")
    
    # Определяем тип поиска
    is_tg_search = query.lower().startswith('tg_')
    search_phone = re.match(r'\+?[0-9\s\-\(\)]{10,15}$', query) or (is_tg_search and re.match(r'\+?[0-9\s\-\(\)]{10,15}$', query[3:]))
    
    print(f"🔍 Это поиск по телефону: {bool(search_phone)}")
    
    # Если это поиск по телефону, используем специальную функцию
    if search_phone and not is_tg_search:
        print("🔍 Используем поиск по телефону")
        return search_all_databases_by_phone(query)
    
    for folder in os.listdir(BASE_DIR):
        db_path = os.path.join(BASE_DIR, folder, 'base.db')
        if not os.path.exists(db_path):
            continue
            
        # Определяем, Telegram ли это база
        is_tg_base = 'telegram' in folder.lower() or 'глаз' in folder.lower()
        
        # Для обычных запросов пропускаем Telegram-базы
        if not is_tg_search and is_tg_base:
            continue
            
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        try:
            if is_tg_search:
                clean_query = query[3:]  # Удаляем префикс tg_
                
                # Если после tg_ идет номер телефона
                if re.match(r'\+?[0-9\s\-\(\)]{10,15}$', clean_query):
                    return search_all_databases_by_phone(clean_query)
                # Иначе ищем по Telegram данным
                elif clean_query.startswith('@'):
                    c.execute("SELECT * FROM people WHERE tg_username LIKE ?", 
                            ('%' + clean_query[1:] + '%',))
                else:
                    c.execute("SELECT * FROM people WHERE tg_id LIKE ?", 
                            ('%' + clean_query + '%',))
            
            elif '@' in query and '.' in query.split('@')[-1]:
                # 🔍 Поиск email в people
                c.execute("SELECT * FROM people WHERE email LIKE ?", ('%' + query + '%',))
                results.extend((folder, row) for row in c.fetchall())
                
                # 🔍 Проверяем, есть ли таблица emails
                c.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='emails'")
                if c.fetchone():
                    c.execute("SELECT * FROM emails WHERE email LIKE ?", ('%' + query + '%',))
                    results.extend((folder, row) for row in c.fetchall())
                continue  # чтобы не дублировать extend ниже
            
            else:
                # Поиск по имени
                c.execute("SELECT * FROM people WHERE name LIKE ?", ('%' + query + '%',))
            
            results.extend((folder, row) for row in c.fetchall())
        
        except Exception as e:
            print(f"❌ Ошибка при поиске в базе {folder}: {e}")
        finally:
            conn.close()
    
    print(f"🔍 Всего найдено результатов: {len(results)}")
    return results


def search_telegram_data(query):
    results = []
    for folder in os.listdir(BASE_DIR):
        # Ищем только в Telegram-папках
        if not ('telegram' in folder.lower() or 'глаз' in folder.lower()):
            continue
            
        db_path = os.path.join(BASE_DIR, folder, 'base.db')
        if not os.path.exists(db_path):
            continue
            
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        if query.startswith('@'):
            c.execute("SELECT * FROM people WHERE tg_username LIKE ?", ('%' + query[1:] + '%',))
        else:
            c.execute("SELECT * FROM people WHERE tg_id LIKE ?", ('%' + query + '%',))
            
        results.extend((folder, row) for row in c.fetchall())
        conn.close()
    
    return results



# --- Telegram Bot Handlers ---

@bot.message_handler(commands=['start'])
def start_message(message):
    markup = telebot.types.InlineKeyboardMarkup()
    markup.add(telebot.types.InlineKeyboardButton("📄 Открыть соглашение", url="https://fenix-osint.site/agreement.html"))
    markup.add(telebot.types.InlineKeyboardButton("☀️ Я ознакомлен", callback_data="agree"))

    agreement_text = (
        "Для использования бота необходимо согласиться с пользовательским соглашением.\n\n"
        "Пожалуйста, ознакомьтесь с ним перед началом работы."
    )

    bot.send_message(
        message.chat.id,
        agreement_text,
        reply_markup=markup
    )


@bot.callback_query_handler(func=lambda call: call.data == "agree")
def handle_agreement_confirm(call):
    welcome_text = "👋 Добро пожаловать в *Fenix-Osint* бот!"
    
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("🔍 Поиск", callback_data="start_search"))
    markup.add(InlineKeyboardButton("☀️ Подписка", callback_data="subscription_menu"))
    markup.add(InlineKeyboardButton("☀️ Баланс", callback_data="balance_menu"))
    markup.add(InlineKeyboardButton("☀️ Информация", callback_data="user_info"))
    
    if is_admin(call.from_user.id):
        markup.add(InlineKeyboardButton("👑 Админ-меню", callback_data="admin_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text=welcome_text,
        parse_mode="Markdown",
        reply_markup=markup
    )
    bot.answer_callback_query(call.id)

@bot.callback_query_handler(func=lambda call: call.data == "admin_menu")
def admin_menu(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("💰 Управление балансом", callback_data="admin_balance_menu"))
    markup.add(InlineKeyboardButton("📅 Управление подписками", callback_data="admin_subs_menu"))
    markup.add(InlineKeyboardButton("📊 Управление статистикой", callback_data="admin_stats_menu"))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="👑 Административное меню:",
        reply_markup=markup
    )
    bot.answer_callback_query(call.id)

# 👇 ВСТАВЬ СЮДА:
@bot.callback_query_handler(func=lambda call: call.data == "main_menu")
def go_main_menu(call):
    try:
        # Сначала отвечаем на callback
        bot.answer_callback_query(call.id)
        
        welcome_text = "👋 Добро пожаловать в Fenix-Osint бот!"
        markup = InlineKeyboardMarkup()
        markup.add(InlineKeyboardButton("🔍 Поиск", callback_data="start_search"))
        markup.add(InlineKeyboardButton("☀️ Подписка", callback_data="subscription_menu"))
        markup.add(InlineKeyboardButton("☀️ Баланс", callback_data="balance_menu"))
        markup.add(InlineKeyboardButton("☀️ Информация", callback_data="user_info"))
        
        if is_admin(call.from_user.id):
            markup.add(InlineKeyboardButton("👑 Управление балансом", callback_data="admin_balance_menu"))
            markup.add(InlineKeyboardButton("👑 Управление подписками", callback_data="admin_subs_menu"))
            markup.add(InlineKeyboardButton("👑 Админ-меню", callback_data="admin_menu"))

        # Получаем текущее сообщение
        try:
            current_text = bot.get_chat(call.message.chat.id).pinned_message.text
        except:
            current_text = None

        # Изменяем только если контент действительно изменился
        if current_text != welcome_text:
            try:
                bot.edit_message_text(
                    chat_id=call.message.chat.id,
                    message_id=call.message.message_id,
                    text=welcome_text,
                    reply_markup=markup
                )
            except Exception as edit_error:
                if "message is not modified" not in str(edit_error):
                    # Если ошибка не связана с неизмененным сообщением
                    bot.send_message(
                        call.message.chat.id,
                        welcome_text,
                        reply_markup=markup
                    )
    except Exception as e:
        print(f"Error in main menu: {e}")
        try:
            bot.send_message(
                call.message.chat.id,
                "Произошла ошибка. Попробуйте еще раз.",
                reply_markup=markup
            )
        except:
            pass

def safe_edit_message(bot, call, new_text, new_markup=None):
    """Безопасное изменение сообщения с проверкой изменений"""
    try:
        # Сначала отвечаем на callback
        bot.answer_callback_query(call.id)
        
        # Получаем текущее сообщение
        try:
            current_msg = bot.get_chat(call.message.chat.id).pinned_message
            current_text = current_msg.text if current_msg else None
            current_markup = current_msg.reply_markup if current_msg else None
        except:
            current_text = None
            current_markup = None

        # Проверяем, нужно ли изменять
        if current_text == new_text and current_markup == new_markup:
            return False

        # Пытаемся изменить
        try:
            bot.edit_message_text(
                chat_id=call.message.chat.id,
                message_id=call.message.message_id,
                text=new_text,
                reply_markup=new_markup
            )
            return True
        except Exception as e:
            if "message is not modified" in str(e):
                return False
            raise e
            
    except Exception as e:
        print(f"Safe edit error: {e}")
        return False

@bot.callback_query_handler(func=lambda call: call.data == "subscription_menu")
def show_subscription_menu(call):
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("☀️ 1 день — 25₽", callback_data="buy_sub_1d"))
    markup.add(InlineKeyboardButton("☀️ 7 дней — 89₽", callback_data="buy_sub_7d"))
    markup.add(InlineKeyboardButton("☀️ 30 дней — 249₽", callback_data="buy_sub_30d"))
    markup.add(InlineKeyboardButton("☀️ 365 дней — 490₽", callback_data="buy_sub_365d"))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="Выберите срок подписки:",
        reply_markup=markup
    )

@bot.callback_query_handler(func=lambda call: call.data.startswith("buy_sub_"))
def buy_subscription(call):
    user_id = call.from_user.id
    options = {
        "buy_sub_1d": (25, "1 день"),
        "buy_sub_7d": (89, "7 дней"),
        "buy_sub_30d": (249, "30 дней"),
        "buy_sub_365d": (490, "365 дней"),
    }
    
    if call.data not in options:
        bot.answer_callback_query(call.id, "❌ Неизвестный тип подписки")
        return

    cost, duration_text = options[call.data]
    balance = get_balance(user_id)
    
    if balance < cost:
        bot.answer_callback_query(call.id, "❌ Недостаточно средств.")
        return

    try:
        bot.answer_callback_query(call.id, f"✅ Покупка подписки на {duration_text} за {cost}₽")
        
        new_text = f"🎉 Подписка на {duration_text} успешно активирована!\n\n💰 Списанно: {cost}₽\n💳 Остаток: {balance - cost}₽"
        markup = InlineKeyboardMarkup()
        markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

        try:
            bot.edit_message_text(
                chat_id=call.message.chat.id,
                message_id=call.message.message_id,
                text=new_text,
                reply_markup=markup
            )
        except:
            bot.send_message(call.message.chat.id, new_text, reply_markup=markup)

        deduct_balance(user_id, cost)
        duration_seconds = {
            "1 день": 86400,
            "7 дней": 7 * 86400,
            "30 дней": 30 * 86400,
            "365 дней": 365 * 86400,
        }[duration_text]
        add_subscription(user_id, duration_seconds)

        # Логируем покупку
        user_info = f"{call.from_user.first_name} {call.from_user.last_name or ''} (@{call.from_user.username or 'нет'})"
        sub_log = f"🎫 ПОКУПКА ПОДПИСКИ\n👤 Пользователь: {user_info}\n🆔 ID: {user_id}\n📦 Подписка: {duration_text}\n💰 Стоимость: {cost}₽"
        send_log_message(sub_log)

    except Exception as e:
        print(f"Ошибка при обработке подписки: {e}")
        bot.answer_callback_query(call.id, "❌ Произошла ошибка при обработке подписки")


@bot.callback_query_handler(func=lambda call: call.data == "balance_menu")
def balance_menu(call):
    user_id = call.from_user.id
    balance = get_balance(user_id)
    
    markup = InlineKeyboardMarkup()
    markup.row(
        InlineKeyboardButton("💳 Пополнить", callback_data="topup_amount"),
        InlineKeyboardButton("📊 История", callback_data="payment_history")
    )
    markup.row(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text=f"💰 Ваш баланс: {balance}₽",
        reply_markup=markup
    )

@bot.callback_query_handler(func=lambda call: call.data == "topup_amount")
def topup_amount(call):
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="balance_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="⚠️ Оплата через бота временно недоступна.\n\n"
             "💳 Для пополнения картой или криптовалютой, обратитесь в техподдержку: @Fenix_Support_BBot",
        reply_markup=markup
    )




user_states = {}  # Храним состояние пользователей
search_states = {}  # Храним состояния поиска


@bot.callback_query_handler(func=lambda call: call.data == "start_search")
def start_search_callback(call):
    user_id = call.from_user.id
    search_states[user_id] = True  # Устанавливаем флаг поиска

    if not is_subscribed(user_id):
        safe_answer_callback_query(bot, call)
        bot.send_message(
            call.message.chat.id,
            "❌ У вас нет активной подписки.\nПожалуйста, приобретите её в разделе 📦 Подписка."
        )
        return

    safe_answer_callback_query(bot, call)

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("❌ Отменить поиск", callback_data="cancel_search"))

    msg = bot.send_message(
        call.message.chat.id,
        "🔍 Введите данные для поиска (ФИО, номер телефона или email):",
        reply_markup=markup
    )
    
    # Сохраняем ID сообщения с кнопкой отмены
    search_states[user_id] = {
        'active': True,
        'cancel_message_id': msg.message_id
    }
    
    bot.register_next_step_handler(msg, process_search_input)

@bot.callback_query_handler(func=lambda call: call.data == "cancel_search")
def cancel_search_handler(call):
    try:
        user_id = call.from_user.id
        
        # Удаляем состояние поиска
        if user_id in search_states:
            del search_states[user_id]
            
        # Удаляем сообщение с запросом ввода
        try:
            bot.delete_message(call.message.chat.id, call.message.message_id)
        except:
            pass
        
        # Просто отвечаем, что поиск отменён
        bot.answer_callback_query(call.id, "❌ Поиск отменён")
    except Exception as e:
        print(f"Error canceling search: {e}")
        bot.answer_callback_query(call.id, "⚠️ Не удалось отменить поиск")



def process_search_input(message: Message):
    user_id = message.from_user.id
    
    # Проверяем, не отменен ли поиск
    if user_id not in search_states or not search_states[user_id].get('active', False):
        return

    if 'cancel_message_id' in search_states[user_id]:
        try:
            bot.delete_message(
                chat_id=message.chat.id,
                message_id=search_states[user_id]['cancel_message_id']
            )
        except:
            pass

    if user_id in search_states:
        del search_states[user_id]

    cleanup_user_messages(user_id, message.chat.id)

    if not is_subscribed(user_id):
        msg = bot.send_message(message.chat.id, "❌ Подписка неактивна.")
        threading.Timer(5.0, lambda: bot.delete_message(msg.chat.id, msg.message_id)).start()
        return

    query = message.text.strip()

    # Логируем начало поиска
    user_info = f"{message.from_user.first_name} {message.from_user.last_name or ''} (@{message.from_user.username or 'нет'})"
    log_text = f"🔍 ПОИСК\n👤 Пользователь: {user_info}\n🆔 ID: {user_id}\n📋 Запрос: {query}"
    send_log_message(log_text)

    # --- Определение типа поиска ---
    if query.startswith(('tg_', 'tg-')):
        search_type = "telegram"
        query = query[3:]
    else:
        is_phone = re.fullmatch(r'\+?[0-9\s\-\(\)]{10,15}', query)
        is_email = '@' in query and '.' in query.split('@')[-1]
        has_name = any(c.isalpha() for c in query)
        dob_pattern = r'\d{2}[./-]\d{2}[./-]\d{4}'
        dob_pattern_iso = r'\d{4}-\d{2}-\d{2}'
        has_dob = re.search(dob_pattern, query) or re.search(dob_pattern_iso, query)
        
        if is_phone:
            search_type = "phone"
        elif is_email:
            search_type = "email"
        elif has_name and has_dob:
            search_type = "name_dob"
        elif has_name and len(query.split()) >= 2:
            search_type = "name"
        else:
            bot.send_message(
                message.chat.id,
                "⚠️ Введите корректные данные для поиска.",
                parse_mode="Markdown"
            )
            return

    bot.send_chat_action(message.chat.id, 'typing')
    search_sticker = bot.send_sticker(
        message.chat.id,
        "CAACAgIAAxkBAAIN9miXYjnEC2M24mONWK2FjHMl0dRWAAIUfQAC2CSxSNb7vkhBE8QnNgQ"
    )
    search_msg = bot.send_message(message.chat.id, "Идёт поиск по нескольким источникам, пожалуйста подождите.")

    try:
        if search_type == "telegram":
            results = search_telegram_data(query)
        elif search_type == "phone":
            results = search_all_databases_by_phone(query)
        else:
            results = search_all_databases_by_query(query)

        try:
            bot.delete_message(message.chat.id, search_sticker.message_id)
            bot.delete_message(message.chat.id, search_msg.message_id)
        except:
            pass

        if not results:
            bot.send_message(message.chat.id, "В наших базах нет информации.")
            result_log = f"❌ РЕЗУЛЬТАТ ПОИСКА\n👤 Пользователь: {user_info}\n📋 Запрос: {query}\n📊 Найдено: 0 результатов"
            send_log_message(result_log)
            return

        for base_name, record in results[:5]:
            report = format_report(record)
            bot.send_message(
                message.chat.id,
                f"📂 База: {base_name}\n\n{report}"
            )

        if len(results) > 5:
            bot.send_message(
                message.chat.id,
                f"ℹ️ Показано 5 из {len(results)} найденных результатов."
            )

        result_log = f"✅ РЕЗУЛЬТАТ ПОИСКА\n👤 Пользователь: {user_info}\n📋 Запрос: {query}\n📊 Найдено: {len(results)} результатов"
        send_log_message(result_log)

    except Exception as e:
        try:
            bot.delete_message(message.chat.id, search_msg.message_id)
        except:
            pass
        bot.send_message(message.chat.id, f"⚠️ Ошибка при поиске: {str(e)}")
        error_log = f"⚠️ ОШИБКА ПОИСКА\n👤 Пользователь: {user_info}\n📋 Запрос: {query}\n❌ Ошибка: {str(e)}"
        send_log_message(error_log)

def search_telegram_data(query):
    """Поиск Telegram данных по ID или username"""
    results = []
    for folder in os.listdir(BASE_DIR):
        # Ищем только в папках с "Telegram" или "глаз" в названии
        if not ('telegram' in folder.lower() or 'глаз' in folder.lower()):
            continue
            
        db_path = os.path.join(BASE_DIR, folder, 'base.db')
        if not os.path.exists(db_path):
            continue

        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        
        # Если запрос начинается с @ - ищем по username
        if query.startswith('@'):
            c.execute("SELECT * FROM people WHERE tg_username LIKE ?", ('%' + query[1:] + '%',))
        else:
            # Иначе ищем по ID
            c.execute("SELECT * FROM people WHERE tg_id LIKE ?", ('%' + query + '%',))
            
        rows = c.fetchall()
        conn.close()
        for row in rows:
            results.append((folder, row))
    
    return results



# Добавляем команду для полной очистки
@bot.message_handler(commands=['clear'])
def clear_chat(message):
    user_id = message.from_user.id
    cleanup_user_messages(user_id, message.chat.id)
    
    # Получаем последние сообщения бота
    messages = bot.get_chat_history(message.chat.id, limit=20)
    for msg in messages:
        if msg.from_user.id == bot.get_me().id:
            try:
                bot.delete_message(message.chat.id, msg.message_id)
            except:
                pass
    
    confirm = bot.send_message(message.chat.id, "✅ Все сообщения очищены")
    threading.Timer(2.0, lambda: bot.delete_message(confirm.chat.id, confirm.message_id)).start()
    try:
        bot.delete_message(message.chat.id, message.message_id)
    except:
        pass


@bot.callback_query_handler(func=lambda call: call.data == "cancel_search")
def cancel_search(call):
    user_id = call.from_user.id
    user_states.pop(user_id, None)
    try:
        bot.delete_message(call.message.chat.id, call.message.message_id)
    except:
        pass
    bot.answer_callback_query(call.id)

@bot.message_handler(func=lambda message: user_states.get(message.from_user.id) == "searching")
def handle_search_query(message):
    user_id = message.from_user.id
    query = message.text.strip()

    bot.send_chat_action(message.chat.id, 'typing')

    # Обработка номера телефона и общего запроса
    phone_match = re.match(r'\+?\d{10,15}$', query)
    
    if phone_match:
        results = search_all_databases_by_phone(query)
    else:
        results = search_all_databases_by_query(query)

    if not results:
        if phone_match:
            operator, region, country = get_phone_info(query)
            report = f"📱 \n├ Телефон: {normalize_phone(query)}"
            if operator:
                report += f"\n├ Оператор: {operator}"
            if region:
                report += f"\n├ Регион: {region}"
            if country:
                report += f"\n└ Страна: {country}"
            report += "\n\n❌ Совпадений не найдено в базах."
            bot.send_message(message.chat.id, report)
        else:
            bot.send_message(message.chat.id, "❌ Совпадений не найдено.")
    else:
        # Ограничим 5 результатами, чтобы не засорять чат
        for base_name, record in results[:5]:
            report = format_report(record)
            bot.send_message(message.chat.id, f"База: {base_name}\n\n{report}")

    user_states.pop(user_id, None)

def init_balance_db():
    os.makedirs('balance_service', exist_ok=True)
    conn = sqlite3.connect('balance_service/balance.db')
    conn.execute('CREATE TABLE IF NOT EXISTS balances (user_id INTEGER PRIMARY KEY, balance INTEGER DEFAULT 0)')
    conn.commit()
    conn.close()


def add_balance(user_id, amount):
    """Пополняет баланс пользователя"""
    try:
        balance_db_path = os.path.join('balance_service', 'balance.db')
        conn = sqlite3.connect(balance_db_path)
        cursor = conn.cursor()
        
        # Создаём таблицу если её нет
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS balances (
                user_id INTEGER PRIMARY KEY,
                balance INTEGER NOT NULL DEFAULT 0
            )
        ''')
        
        cursor.execute('''
            INSERT OR IGNORE INTO balances (user_id, balance)
            VALUES (?, 0)
        ''', (user_id,))
        
        cursor.execute('''
            UPDATE balances
            SET balance = balance + ?
            WHERE user_id = ?
        ''', (amount, user_id))
        
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"[Ошибка пополнения баланса]: {e}")
        return False

def deduct_balance(user_id, amount):
    conn = sqlite3.connect('balance_service/balance.db')
    c = conn.cursor()
    c.execute('UPDATE balances SET balance = balance - ? WHERE user_id = ? AND balance >= ?', (amount, user_id, amount))
    conn.commit()
    conn.close()

def get_balance(user_id):
    conn = sqlite3.connect('balance_service/balance.db')
    cursor = conn.cursor()

    # создаём таблицу, если её нет
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS balances (
            user_id INTEGER PRIMARY KEY,
            balance INTEGER NOT NULL DEFAULT 0
        )
    ''')

    cursor.execute("SELECT balance FROM balances WHERE user_id = ?", (user_id,))
    result = cursor.fetchone()
    conn.close()

    return result[0] if result else 0



def is_subscribed(user_id):
    """Проверяет активна ли подписка у пользователя"""
    try:
        # Убедимся, что база данных существует
        init_subscription_db()
            
        subs_db_path = os.path.join('subscriptions_service', 'subs.db')
        conn = sqlite3.connect(subs_db_path)
        cursor = conn.cursor()
            
        cursor.execute("SELECT expires_at FROM subscriptions WHERE user_id = ?", (user_id,))
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return time.time() < result[0]
        return False
    except Exception as e:
        print(f"[Ошибка проверки подписки]: {e}")
        return False

def add_subscription(user_id, duration_seconds):
    """Добавляет или обновляет подписку пользователя"""
    try:
        # Убедимся, что база данных существует
        init_subscription_db()
        
        # Проверка максимального срока (5000 дней)
        max_days = 5000
        max_seconds = max_days * 86400
        
        if duration_seconds > max_seconds:
            duration_seconds = max_seconds
            print(f"⚠️ Срок подписки ограничен {max_days} днями")
        
        subs_db_path = os.path.join('subscriptions_service', 'subs.db')
        conn = sqlite3.connect(subs_db_path)
        cursor = conn.cursor()
        
        expires_at = int(time.time()) + duration_seconds
        cursor.execute('''
            INSERT OR REPLACE INTO subscriptions (user_id, expires_at)
            VALUES (?, ?)
        ''', (user_id, expires_at))
        
        conn.commit()
        conn.close()
        print(f"✅ Подписка добавлена для пользователя {user_id}, истекает: {datetime.fromtimestamp(expires_at)}")
        return True
    except Exception as e:
        print(f"[Ошибка добавления подписки]: {e}")
        return False


def init_subscription_db():
    """Инициализация базы данных подписок"""
    os.makedirs('subscriptions_service', exist_ok=True)
    conn = sqlite3.connect('subscriptions_service/subs.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS subscriptions (
            user_id INTEGER PRIMARY KEY,
            expires_at INTEGER NOT NULL
        )
    ''')
    conn.commit()
    conn.close()
    print("✅ База данных подписок инициализирована")


# Добавьте этот вызов в начало main_menu() или в init_databases()
def init_databases():
    """Инициализирует все необходимые базы данных"""
    try:
        # Инициализация базы подписок
        init_subscription_db()
        
        # Остальной код инициализации...
        os.makedirs('balance_service', exist_ok=True)
        balance_db_path = os.path.join('balance_service', 'balance.db')
        balance_conn = sqlite3.connect(balance_db_path)
        balance_cursor = balance_conn.cursor()
        
        # Таблица балансов
        balance_cursor.execute('''
            CREATE TABLE IF NOT EXISTS balances (
                user_id INTEGER PRIMARY KEY,
                balance INTEGER NOT NULL DEFAULT 0
            )
        ''')
        
        # Таблица платежей
        balance_cursor.execute('''
            CREATE TABLE IF NOT EXISTS payments (
                payment_id TEXT PRIMARY KEY,
                user_id INTEGER,
                amount INTEGER,
                status TEXT DEFAULT "pending",
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        balance_conn.commit()
        balance_conn.close()
        
        # База данных для статистики
        stats_conn = sqlite3.connect('stats.db')
        stats_cursor = stats_conn.cursor()
        stats_cursor.execute('''
            CREATE TABLE IF NOT EXISTS stats (
                base_name TEXT PRIMARY KEY,
                record_count INTEGER
            )
        ''')
        
        # Проверяем, есть ли данные в таблице stats
        stats_cursor.execute("SELECT COUNT(*) FROM stats")
        if stats_cursor.fetchone()[0] == 0:
            # Добавляем тестовые данные, если таблица пуста
            stats_cursor.execute("INSERT INTO stats (base_name, record_count) VALUES (?, ?)", 
                               ("Основная база", 1000000))
        
        stats_conn.commit()
        stats_conn.close()
        
        print("✅ Все базы данных успешно инициализированы")
    except Exception as e:
        print(f"❌ Критическая ошибка инициализации баз данных: {e}")
        sys.exit(1)

def init_payments_db():
    """Инициализация таблицы платежей"""
    conn = sqlite3.connect('balance_service/balance.db')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS payments (
            payment_id TEXT PRIMARY KEY,
            user_id INTEGER,
            amount INTEGER,
            status TEXT DEFAULT "pending",
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()


def send_payment_notification(user_id, amount):
    """Отправляет уведомление о successful payment"""
    try:
        bot.send_message(
            user_id,
            f"✅ Ваш баланс пополнен на {amount}₽!\n"
            f"Новый баланс: {get_balance(user_id)}₽"
        )
    except Exception as e:
        print(f"Ошибка отправки уведомления: {e}")


ALLOWED_GROUP_ID = -1002991202021 # ID разрешённой группы
PHOTO_PATH = "startup.jpg"         # Путь к картинке для старта


def leave_unallowed_chats():
    try:
        updates = bot.get_updates(limit=100)
        for update in updates:
            msg = update.message
            if msg and msg.chat.type in ["group", "supergroup"]:
                if msg.chat.id != ALLOWED_GROUP_ID:
                    try:
                        bot.leave_chat(msg.chat.id)
                    except:
                        pass
    except:
        pass


def send_startup_message():
    try:
        total_records = get_total_records_count()
        start_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        text = (
            f"✅ Бот запущен!\n"
            f"⏰ Время запуска: {start_time}(+3 часа)\n"
            f"📂 Всего записей во всех базах: {total_records:,}"
        )
        with open(PHOTO_PATH, "rb") as photo:
            bot.send_photo(ALLOWED_GROUP_ID, photo, caption=text)

        # Логирование запуска
        send_log_message("🚀 Бот запущен и готов к работе!")

    except Exception as e:
        print(f"Не удалось отправить сообщение о запуске: {e}")



def main_menu():
    """Главное меню в терминале"""
    init_databases()
    while True:
        print("\n" + "="*20 + " FENIX OSINT " + "="*20)
        print("1. Импорт базы данных")
        print("2. Запустить Telegram-бота")
        print("3. Выйти")

        choice = input("\nВыберите действие (1-3): ").strip()

        if choice == '1':
            import_database()
        elif choice == '2':
            leave_unallowed_chats()
            send_startup_message()
            run_bot()
        elif choice == '3':
            print("Завершение работы...")
            break
        else:
            print("❌ Неверный выбор. Введите число от 1 до 3")

    


@bot.callback_query_handler(func=lambda call: call.data == "admin_balance_menu")
def admin_balance_menu(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("➕ Пополнить баланс", callback_data="admin_add_balance"))
    markup.add(InlineKeyboardButton("🔄 Обнулить баланс", callback_data="admin_reset_balance"))
    markup.add(InlineKeyboardButton("🗑 Удалить запись", callback_data="admin_delete_record"))  # Новая кнопка
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="👑 Административное меню управления балансами:",
        reply_markup=markup
    )

@bot.callback_query_handler(func=lambda call: call.data == "admin_delete_record")
def admin_delete_record(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите данные для удаления (ФИО или номер телефона):"
    )
    bot.register_next_step_handler(msg, process_admin_delete_record_search)

def process_admin_delete_record_search(message):
    query = message.text.strip()
    results = search_all_databases_by_query(query)
    
    if not results:
        bot.send_message(message.chat.id, "В наших базах нет информации.")
        return

    # Ограничим вывод 10 результатами
    for i, (base_name, record) in enumerate(results[:10]):
        _, name, dob, phone, email, address = record
        text = f"🔍 Результат {i+1} (База: {base_name}):\n"
        if name: text += f"👤 {name}\n"
        if dob: text += f"🎂 {dob}\n"
        if phone: text += f"📞 {phone}\n"
        if email: text += f"📧 {email}\n"
        
        markup = InlineKeyboardMarkup()
        markup.add(InlineKeyboardButton(
            "🗑 Удалить эту запись", 
            callback_data=f"confirm_delete_{base_name}_{record[0]}"  # ID записи
        ))
        
        bot.send_message(
            message.chat.id,
            text,
            reply_markup=markup
        )

@bot.callback_query_handler(func=lambda call: call.data.startswith("confirm_delete_"))
def confirm_delete_record(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    parts = call.data.split('_')
    base_name = parts[2]
    record_id = parts[3]
    
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("✅ Да, удалить", callback_data=f"final_delete_{base_name}_{record_id}"))
    markup.add(InlineKeyboardButton("❌ Нет, отменить", callback_data="cancel_delete"))
    
    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="⚠️ Вы уверены, что хотите удалить эту запись?",
        reply_markup=markup
    )

@bot.callback_query_handler(func=lambda call: call.data.startswith("final_delete_"))
def final_delete_record(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    parts = call.data.split('_')
    base_name = parts[2]
    record_id = parts[3]
    
    db_path = os.path.join(BASE_DIR, base_name, 'base.db')
    if not os.path.exists(db_path):
        bot.answer_callback_query(call.id, "❌ База данных не найдена!")
        return

    try:
        conn = sqlite3.connect(db_path)
        c = conn.cursor()
        c.execute("DELETE FROM people WHERE id = ?", (record_id,))
        conn.commit()
        conn.close()
        
        bot.edit_message_text(
            chat_id=call.message.chat.id,
            message_id=call.message.message_id,
            text="✅ Запись успешно удалена!"
        )
    except Exception as e:
        bot.answer_callback_query(call.id, f"❌ Ошибка: {str(e)}")

@bot.callback_query_handler(func=lambda call: call.data == "cancel_delete")
def cancel_delete_record(call):
    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="❌ Удаление отменено."
    )

# Обработчик пополнения баланса
@bot.callback_query_handler(func=lambda call: call.data == "admin_add_balance")
def admin_add_balance(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите ID пользователя и сумму для пополнения через пробел (например: 123456789 100):"
    )
    bot.register_next_step_handler(msg, process_admin_add_balance)

def process_admin_add_balance(message):
    try:
        parts = message.text.strip().split()
        if len(parts) != 2:
            bot.send_message(message.chat.id, "❌ Неверный формат. Используйте: <ID> <сумма>")
            return

        user_id, amount = map(int, parts)

        if add_balance(user_id, amount):
            new_balance = get_balance(user_id)

            bot.send_message(
                message.chat.id,
                f"✅ Баланс пользователя {user_id} увеличен на {amount}₽\n"
                f"💰 Новый баланс: {new_balance}₽"
            )

            try:
                markup = InlineKeyboardMarkup()
                markup.add(InlineKeyboardButton("Скрыть", callback_data="hide_msg"))
                bot.send_message(
                    user_id,
                    f"✅ Ваш баланс пополнен на {amount}₽!\n💰 Новый баланс: {new_balance}₽",
                    reply_markup=markup
                )
            except:
                bot.send_message(message.chat.id, f"⚠️ Не удалось уведомить пользователя {user_id}")

            # Логирование
            admin_info = f"{message.from_user.first_name} {message.from_user.last_name or ''}"
            balance_log = f"💰 ПОПОЛНЕНИЕ БАЛАНСА\n👑 Админ: {admin_info}\n👤 Пользователь: ID {user_id}\n💵 Сумма: {amount}₽\n💳 Новый баланс: {new_balance}₽"
            send_log_message(balance_log)

        else:
            bot.send_message(message.chat.id, "❌ Ошибка при пополнении баланса")

    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Ошибка: {e}\nФормат: <ID> <сумма>")



# обработчик кнопки "❌ Скрыть"
@bot.callback_query_handler(func=lambda call: call.data == "hide_msg")
def hide_message(call):
    try:
        bot.delete_message(call.message.chat.id, call.message.message_id)
    except:
        bot.answer_callback_query(call.id, "⚠️ Не удалось удалить сообщение")





# Обработчик обнуления баланса
@bot.callback_query_handler(func=lambda call: call.data == "admin_reset_balance")
def admin_reset_balance(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите ID пользователя для обнуления баланса:"
    )
    bot.register_next_step_handler(msg, process_admin_reset_balance)

def process_admin_reset_balance(message):
    try:
        user_id = int(message.text)
        conn = sqlite3.connect('balance_service/balance.db')
        c = conn.cursor()
        c.execute('UPDATE balances SET balance = 0 WHERE user_id = ?', (user_id,))
        conn.commit()
        conn.close()
        bot.send_message(
            message.chat.id,
            f"✅ Баланс пользователя {user_id} обнулён."
        )
    except Exception as e:
        bot.send_message(
            message.chat.id,
            f"❌ Ошибка: {e}\nПожалуйста, введите корректный ID пользователя"
        )

def safe_answer_callback_query(bot, call, text=None, show_alert=False):
    """
    Безопасная обработка callback-запросов с проверкой на устаревание
    """
    try:
        bot.answer_callback_query(call.id, text=text, show_alert=show_alert)
    except ApiTelegramException as e:
        if "query is too old" in str(e):
            print(f"Callback query expired: {e}")
        else:
            raise e

@bot.callback_query_handler(func=lambda call: call.data == "admin_subs_menu")
def admin_subs_menu(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ Доступ запрещен!")
        return

    menu_text = "👑 Админ: Управление подписками"
    markup = InlineKeyboardMarkup(row_width=2)
    markup.add(
        InlineKeyboardButton("➕ Добавить подписку", callback_data="admin_add_sub"),
        InlineKeyboardButton("❌ Удалить подписку", callback_data="admin_remove_sub"),
        InlineKeyboardButton("📊 Проверить подписку", callback_data="admin_check_sub"),
        InlineKeyboardButton("❮ Назад", callback_data="main_menu")
    )

    if not safe_edit_message(bot, call, menu_text, markup):
        # Если не удалось изменить, отправляем новое сообщение
        bot.send_message(
            call.message.chat.id,
            menu_text,
            reply_markup=markup
        )
    bot.answer_callback_query(call.id)

@bot.callback_query_handler(func=lambda call: call.data == "admin_add_sub")
def admin_add_sub(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ Доступ запрещен!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите ID пользователя и срок подписки в днях через пробел:\nПример: 123456789 30"
    )
    bot.register_next_step_handler(msg, process_admin_add_sub)

def process_admin_add_sub(message):
    try:
        user_id, days = map(int, message.text.split())
        expires_at = int(time.time()) + days * 86400
        
        with db_lock:
            conn = sqlite3.connect('subscriptions_service/subs.db')
            cursor = conn.cursor()
            cursor.execute(
                "INSERT OR REPLACE INTO subscriptions (user_id, expires_at) VALUES (?, ?)",
                (user_id, expires_at)
            )
            conn.commit()
            conn.close()

        bot.send_message(
            message.chat.id,
            f"✅ Пользователю {user_id} добавлена подписка на {days} дней"
        )
    except Exception as e:
        bot.send_message(
            message.chat.id,
            f"❌ Ошибка: {e}\nФормат: ID_пользователя КОЛИЧЕСТВО_ДНЕЙ"
        )

@bot.callback_query_handler(func=lambda call: call.data == "admin_check_sub")
def admin_check_sub(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ Доступ запрещен!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите ID пользователя для проверки подписки:"
    )
    bot.register_next_step_handler(msg, process_admin_check_sub)

def process_admin_check_sub(message):
    try:
        user_id = int(message.text)
        subscribed = is_subscribed(user_id)
        
        if subscribed:
            with db_lock:
                conn = sqlite3.connect('subscriptions_service/subs.db')
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT expires_at FROM subscriptions WHERE user_id = ?",
                    (user_id,)
                )
                expires_at = cursor.fetchone()[0]
                conn.close()
            
            remaining = (expires_at - time.time()) / 86400
            text = (
                f"✅ Пользователь {user_id} имеет активную подписку\n"
                f"⌛ Осталось: {remaining:.1f} дней\n"
                f"📅 Дата окончания: {datetime.fromtimestamp(expires_at)}"
            )
        else:
            text = f"❌ Пользователь {user_id} не имеет активной подписки"

        bot.send_message(message.chat.id, text)
    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Ошибка: {e}")

@bot.callback_query_handler(func=lambda call: call.data == "admin_remove_sub")
def admin_remove_sub(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ Доступ запрещен!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите ID пользователя для удаления подписки:"
    )
    bot.register_next_step_handler(msg, process_admin_remove_sub)

def process_admin_remove_sub(message):
    try:
        user_id = int(message.text)
        
        with db_lock:
            conn = sqlite3.connect('subscriptions_service/subs.db')
            cursor = conn.cursor()
            cursor.execute(
                "DELETE FROM subscriptions WHERE user_id = ?",
                (user_id,)
            )
            conn.commit()
            conn.close()

        bot.send_message(
            message.chat.id,
            f"✅ Подписка пользователя {user_id} удалена"
        )
    except Exception as e:
        bot.send_message(
            message.chat.id,
            f"❌ Ошибка: {e}"
        )

def safe_db_execute(func):
    """Декоратор для безопасного выполнения запросов к БД"""
    def wrapper(*args, **kwargs):
        try:
            with db_lock:
                return func(*args, **kwargs)
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e):
                time.sleep(0.5)
                return wrapper(*args, **kwargs)
            raise e
    return wrapper

db_lock = threading.Lock()

def run_bot():
    """Запуск Telegram бота"""
    try:
        # Гарантированно пересоздаём все БД
        init_databases()        # вызывает init_subscription_db и init_balance_db
        print("🤖 Бот запущен...")
        bot.infinity_polling()
    except Exception as e:
        print(f"🚨 Ошибка бота: {e}")
        print("♻️ Перезапуск через 5 секунд...")
        time.sleep(5)
        run_bot()




@bot.callback_query_handler(func=lambda call: call.data == "top_up_balance")
async def top_up_balance(call):
    user_id = call.from_user.id
    markup = InlineKeyboardMarkup()
    
    # Варианты сумм для пополнения
    amounts = [15, 25, 50, 100, 500, 1000]
    for amount in amounts:
        markup.add(InlineKeyboardButton(
            f"{amount}₽", 
            callback_data=f"create_payment_{amount}"
        ))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="balance_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="💳 Выберите сумму пополнения:",
        reply_markup=markup
    )






@bot.callback_query_handler(func=lambda call: call.data.startswith("check_payment_"))
def check_payment(call):
    try:
        # Отвечаем сразу на callback
        try:
            bot.answer_callback_query(call.id)
        except:
            pass
        
        payment_id = call.data.split("_")[-1]
        user_id = call.from_user.id
        
        print(f"🔍 Проверка платежа: {payment_id}")
        
        # Пробуем сначала локальную проверку
        result = check_payment_local(payment_id)
        
        # Если локальная проверка не нашла, пробуем удаленную
        if result.get("status") == "not_found":
            print("⚠️ Локально не найден, пробуем удаленную проверку")
            result = check_payment_remote(payment_id)
        
        status = result.get("status")
        print(f"📊 Статус платежа {payment_id}: {status}")
        
        if status == "completed":
            amount = result.get("amount", 0)
            balance = get_balance(user_id)
            
            bot.send_message(
                user_id,
                f"✅ Оплата подтверждена!\n💰 Зачислено: {amount}₽\n💳 Текущий баланс: {balance}₽"
            )
            
        elif status == "pending":
            bot.send_message(
                user_id,
                "⌛ Ожидаем подтверждения оплаты\n\n" +
                "Обычно это занимает 1-2 минуты. Попробуйте проверить через минуту."
            )
            
        elif status == "expired":
            bot.send_message(
                user_id,
                "❌ Платёж просрочен\n\n" +
                "Ссылка для оплаты действительна 2 часа. Создайте новый платёж."
            )
            
        elif status == "not_found":
            bot.send_message(user_id, "❌ Платёж не найден в системе")
            
        else:
            bot.send_message(
                user_id,
                "⚠️ Временные технические проблемы\n\n" +
                "Попробуйте проверить позже или обратитесь в поддержку @Fenix_Support_BBot"
            )
            
    except Exception as e:
        print(f"❌ Ошибка в обработчике платежа: {e}")
        try:
            bot.send_message(
                call.from_user.id,
                "⚠️ Временные технические проблемы\n\n" +
                "Попробуйте проверить позже или обратитесь в поддержку @Fenix_Support_BBot"
            )
        except:
            pass



# --- Функция для сохранения платежа в БД ---
def save_payment_to_db(user_id, payment_id, amount):
    """Сохраняет платёж в БД"""
    try:
        conn = sqlite3.connect('balance_service/balance.db')
        cursor = conn.cursor()
        
        # Создаем таблицу если её нет
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS payments (
                payment_id TEXT PRIMARY KEY,
                user_id INTEGER,
                amount INTEGER,
                status TEXT DEFAULT "pending",
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute(
            "INSERT OR REPLACE INTO payments (payment_id, user_id, amount, status) VALUES (?, ?, ?, ?)",
            (payment_id, user_id, amount, 'pending')
        )
        conn.commit()
        conn.close()
        
        print(f"✅ Платёж сохранен в БД: {payment_id}, user: {user_id}, amount: {amount}")
        return True
        
    except Exception as e:
        print(f"❌ Ошибка сохранения платежа в БД: {e}")
        return False




# Глобальная блокировка для работы с БД
db_lock = threading.Lock()

WEBHOOK_SECRET = "https://fenix-webhook.site/hook.php"

def verify_signature(payload, signature):
    computed_signature = hmac.new(
        WEBHOOK_SECRET.encode(),
        payload,
        sha256
    ).hexdigest()
    return hmac.compare_digest(computed_signature, signature)




@bot.callback_query_handler(func=lambda call: call.data == "payment_history")
def payment_history(call):
    user_id = call.from_user.id
    conn = sqlite3.connect('balance_service/balance.db')
    cursor = conn.cursor()
    
    cursor.execute(
        "SELECT amount, status, created_at FROM payments WHERE user_id = ? ORDER BY created_at DESC LIMIT 5",
        (user_id,)
    )
    payments = cursor.fetchall()
    conn.close()
    
    if not payments:
        text = "📭 История платежей пуста"
    else:
        text = "📜 Последние 5 платежей:\n\n"
        for amount, status, date in payments:
            status_icon = "✅" if status == "completed" else "⌛"
            text += f"{status_icon} {amount}₽ - {date}\n"
    
    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="balance_menu"))
    
    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text=text,
        reply_markup=markup
    )






def create_payment_direct(amount, user_id):
    url = "https://api.memepay.io/v1/payments"
    headers = {
        "Authorization": f"Bearer {MEMEPAY_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "amount": amount,
        "shop_id": MEMEPAY_SHOP_ID,
        "currency": "RUB",
        "metadata": {"user_id": user_id}
    }
    
    try:
        response = requests.post(url, json=data, headers=headers, timeout=10)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        print(f"Ошибка запроса: {e}\nОтвет сервера: {e.response.text if hasattr(e, 'response') else 'нет ответа'}")
        raise

def create_yoomoney_payment(user_id, amount):
    import uuid
    payment_id = str(uuid.uuid4())
    save_payment_to_db(user_id, payment_id, amount)  # в твоей БД уже есть функция

    params = {
        "receiver": YOOMONEY_RECEIVER,
        "quickpay-form": "shop",
        "targets": f"Пополнение баланса для {user_id}",
        "paymentType": "AC",   # AC = карта, PC = кошелёк
        "sum": amount,
        "label": payment_id,
    }
    url = "https://yoomoney.ru/quickpay/confirm.xml?" + urlencode(params)
    return payment_id, url



@bot.callback_query_handler(func=lambda call: call.data.startswith("create_payment_select_"))
def handle_payment_creation(call):
    user_id = call.from_user.id
    amount = int(call.data.split("_")[-1])

    payment_id = str(uuid.uuid4())
    save_payment_to_db(user_id, payment_id, amount)

    # YooMoney ссылка
    params = {
        "receiver": YOOMONEY_RECEIVER,
        "quickpay-form": "shop",
        "targets": f"Пополнение баланса для {user_id}",
        "paymentType": "AC",
        "sum": amount,
        "label": payment_id,
    }
    url = "https://yoomoney.ru/quickpay/confirm.xml?" + urlencode(params)

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("💳 Оплатить", url=url))
    markup.add(InlineKeyboardButton("🔄 Проверить оплату", callback_data=f"check_payment_{payment_id}"))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="balance_menu"))

    current_time = datetime.now().strftime("%H:%M:%S")
    
    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text=f"💰 Сумма: {amount}₽\n⏰ Создано: {current_time}\n\n" +
             "После оплаты нажмите «Проверить оплату».\n" +
             "⚠️ Ссылка действительна 2 часа",
        reply_markup=markup
    )




@bot.callback_query_handler(func=lambda call: True)
def handle_callback(call):  # Обработчик callback-запросов
    if call.data == "user_info":
        try:
            user = call.from_user
            full_name = f"{user.first_name} {user.last_name}" if user.last_name else user.first_name
            sub_info = get_subscription_info(user.id)
            total_records = get_manual_stats_count()
            current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            
            message_text = (
                f"🆔 Ваш ID: {user.id}\n"
                f"👤 Имя: {full_name}\n"
                f"🕒 Текущее время: {current_time}\n"
                f"⏳ Статус подписки: {sub_info['status']}\n"
                f"📅 {sub_info['expiry_info']}\n\n"
                f"🛠 Тех-Поддержка: @Fenix_Support_BBot\n"
                f"🛠 Наш сайт: FENIX-OSINT.SITE\n"
            )
                
            markup = InlineKeyboardMarkup()
            markup.row(
                InlineKeyboardButton("❮ Назад", callback_data="main_menu"),
                InlineKeyboardButton("Удалить себя", url="https://fenix-osint.site/removal.html"),
                InlineKeyboardButton("Базы данных", url="https://fenix-osint.site/base.html")   
            )
            
            bot.edit_message_text(
                chat_id=call.message.chat.id,
                message_id=call.message.message_id,
                text=message_text,
                reply_markup=markup
            )
            bot.answer_callback_query(call.id)
        except Exception as e:
            print(f"Error in user_info: {e}")
            bot.answer_callback_query(
                call.id,
                text="❌ Произошла ошибка при получении информации",
                show_alert=True
            )
    elif call.data == "main_menu":
        # Обработка возврата в главное меню
        go_main_menu(call)

def get_manual_stats_count():
    """Получает общее количество записей из ручной статистики"""
    try:
        ensure_stats_table_exists()  # Добавьте эту строку
        conn = sqlite3.connect('stats.db')
        c = conn.cursor()
        c.execute("SELECT SUM(record_count) FROM stats")
        total = c.fetchone()[0] or 0
        conn.close()
        return total
    except Exception as e:
        print(f"Error getting manual stats: {e}")
        return 0



def get_subscription_info(user_id):
    """Получает информацию о подписке пользователя"""
    try:
        conn = sqlite3.connect('subscriptions_service/subs.db')
        cursor = conn.cursor()
        cursor.execute(
            "SELECT expires_at FROM subscriptions WHERE user_id = ?",
            (user_id,)
        )
        result = cursor.fetchone()
        conn.close()
        
        if result:
            expires_at = result[0]
            current_time = time.time()
            if current_time < expires_at:
                remaining = expires_at - current_time
                days = int(remaining // 86400)
                hours = int((remaining % 86400) // 3600)
                minutes = int((remaining % 3600) // 60)
                return {
                    'status': "✅ Активна",
                    'expiry_info': f"Осталось: {days} дн. {hours} час. {minutes} мин."
                }
            else:
                return {
                    'status': "❌ Истекла",
                    'expiry_info': "Подписка истекла"
                }
        
        return {
            'status': "❌ Неактивна",
            'expiry_info': "Подписка отсутствует"
        }
    except Exception as e:
        print(f"Ошибка получения информации о подписке: {e}")
        return {
            'status': "❌ Ошибка проверки",
            'expiry_info': "Не удалось проверить статус подписки"
        }

def get_total_records_count():
    """Подсчитывает общее количество записей во всех базах"""
    total = 0
    for folder in os.listdir(BASE_DIR):
        db_path = os.path.join(BASE_DIR, folder, 'base.db')
        if os.path.exists(db_path):
            try:
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                cursor.execute("SELECT COUNT(*) FROM people")
                count = cursor.fetchone()[0]
                total += count
                conn.close()
            except:
                continue
    return total




@bot.callback_query_handler(func=lambda call: call.data == "admin_stats_menu")
def admin_stats_menu(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    markup = InlineKeyboardMarkup()
    markup.add(InlineKeyboardButton("✏️ Изменить статистику", callback_data="edit_stats"))
    markup.add(InlineKeyboardButton("📊 Показать статистику", callback_data="show_stats"))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))

    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text="📊 Меню управления статистикой:",
        reply_markup=markup
    )

@bot.callback_query_handler(func=lambda call: call.data == "edit_stats")
def edit_stats(call):
    if not is_admin(call.from_user.id):
        bot.answer_callback_query(call.id, "❌ У вас нет прав доступа!")
        return

    msg = bot.send_message(
        call.message.chat.id,
        "Введите название базы и количество записей в формате:\n"
        "`Название базы:количество`\n\n"
        "Пример: `Россия:1500000`"
    )
    bot.register_next_step_handler(msg, process_stats_update)

def process_stats_update(message):
    try:
        ensure_stats_table_exists()
        base_name, count = message.text.split(':')
        count = int(count.strip())
        base_name = base_name.strip()
        
        conn = sqlite3.connect('stats.db')
        c = conn.cursor()
        c.execute(
            "INSERT OR REPLACE INTO stats (base_name, record_count) VALUES (?, ?)",
            (base_name, count)
        )
        conn.commit()
        conn.close()
        
        bot.send_message(
            message.chat.id,
            f"✅ Статистика для базы '{base_name}' обновлена: {count:,} записей"
        )
    except Exception as e:
        bot.send_message(
            message.chat.id,
            f"❌ Ошибка: {e}\nИспользуйте формат: Название базы:количество"
        )

@bot.callback_query_handler(func=lambda call: call.data == "show_stats")
def show_stats(call):
    ensure_stats_table_exists()
    conn = sqlite3.connect('stats.db')
    c = conn.cursor()
    c.execute("SELECT base_name, record_count FROM stats ORDER BY base_name")
    stats = c.fetchall()
    conn.close()
    
    if not stats:
        text = "📊 Статистика по базам:\n\nНет данных"
    else:
        text = "📊 Статистика по базам:\n\n"
        total = 0
        for base_name, count in stats:
            text += f"▪ {base_name}: {count:,}\n"
            total += count
        text += f"\n📌 Всего: {total:,} записей"
    
    markup = InlineKeyboardMarkup()
    if is_admin(call.from_user.id):
        markup.add(InlineKeyboardButton("✏️ Изменить", callback_data="edit_stats"))
    markup.add(InlineKeyboardButton("❮ Назад", callback_data="main_menu"))
    
    bot.edit_message_text(
        chat_id=call.message.chat.id,
        message_id=call.message.message_id,
        text=text,
        reply_markup=markup
    )



@bot.message_handler(commands=['statistic'])
def send_statistic(message):  # Здесь message - параметр функции
    try:
        user = message.from_user
        full_name = f"{user.first_name} {user.last_name}" if user.last_name else user.first_name
        sub_info = get_subscription_info(user.id)
        total_records = get_manual_stats_count()
        current_time = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        
        message_text = (
                f"🆔 Ваш ID: {user.id}\n"
                f"👤 Имя: {full_name}\n"
                f"🕒 Текущее время: {current_time}\n"
                f"⏳ Статус подписки: {sub_info['status']}\n"
                f"📅 {sub_info['expiry_info']}\n\n"
                f"🛠 Тех-Поддержка: @Fenix_Support_BBot\n"
                f"🛠 Наш сайт: FENIX-OSINT.SITE\n"
            )
        
        markup = InlineKeyboardMarkup()
        markup.row(
            InlineKeyboardButton("❮ Назад", callback_data="main_menu"),
            InlineKeyboardButton("Удалить себя", url="https://fenix-osint.site"),
            InlineKeyboardButton("Базы данных", url="https://fenix-osint.site/base.html")   
        )
        
        bot.send_message(
            chat_id=message.chat.id,  # Используем message.chat.id
            text=message_text,
            reply_markup=markup
        )
    except Exception as e:
        print(f"Error in /statistic command: {e}")
        bot.send_message(
            chat_id=message.chat.id,
            text="❌ Произошла ошибка при получении статистики"
        )



user_messages = {}

@bot.message_handler(func=lambda message: True)
def handle_all_messages(message):
    # Сохраняем ID сообщения пользователя
    if message.from_user.id not in user_messages:
        user_messages[message.from_user.id] = []
    user_messages[message.from_user.id].append(message.message_id)
    
    # Обрабатываем команды
    if message.text.startswith('/'):
        bot.process_new_messages([message])
        return

    # Удаляем сообщение пользователя через 1 секунду (чтобы успеть обработать)
    threading.Timer(1.0, delete_user_message, args=[message]).start()

def delete_user_message(message):
    try:
        bot.delete_message(message.chat.id, message.message_id)
        # Удаляем из нашего списка
        if message.from_user.id in user_messages and message.message_id in user_messages[message.from_user.id]:
            user_messages[message.from_user.id].remove(message.message_id)
    except Exception as e:
        print(f"Не удалось удалить сообщение: {e}")

def cleanup_user_messages(user_id, chat_id):
    """Очистка всех сообщений пользователя"""
    if user_id in user_messages:
        for msg_id in user_messages[user_id][:]:  # Делаем копию списка
            try:
                bot.delete_message(chat_id, msg_id)
                user_messages[user_id].remove(msg_id)
            except Exception as e:
                print(f"Не удалось удалить сообщение {msg_id}: {e}")
 

def get_balance(user_id: int) -> int:
    """Локальный баланс в файле balance_service/balance.db"""
    try:
        os.makedirs('balance_service', exist_ok=True)
        conn = sqlite3.connect('balance_service/balance.db')
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS balances (
                user_id INTEGER PRIMARY KEY,
                balance INTEGER NOT NULL DEFAULT 0
            )
        ''')
        # гарантируем, что запись для пользователя есть
        c.execute('INSERT OR IGNORE INTO balances (user_id, balance) VALUES (?, 0)', (user_id,))
        c.execute('SELECT balance FROM balances WHERE user_id = ?', (user_id,))
        row = c.fetchone()
        return int(row[0]) if row else 0
    except Exception as e:
        print(f"[Ошибка получения баланса]: {e}")
        return 0
    finally:
        try:
            conn.close()
        except:
            pass




@bot.message_handler(commands=["addbalance"])
def cmd_addbalance(message):
    try:
        parts = message.text.split()
        if len(parts) != 3:
            bot.send_message(message.chat.id, "❌ Формат: /addbalance <user_id> <сумма>")
            return

        _, uid, amt = parts
        uid, amt = int(uid), int(amt)

        if add_balance(uid, amt):
            new_balance = get_balance(uid)

            # Сообщение админу
            bot.send_message(
                message.chat.id,
                f"✅ Баланс пользователя {uid} увеличен на {amt}₽\n💰 Новый баланс: {new_balance}₽"
            )

            # Сообщение пользователю
            try:
                markup = InlineKeyboardMarkup()
                markup.add(InlineKeyboardButton("Скрыть", callback_data="hide_msg"))
                bot.send_message(
                    uid,
                    f"✅ Ваш баланс пополнен администратором на сумму - {amt}₽\n"
                    f"💳 Ваш текущий баланс - {new_balance}₽",
                    reply_markup=markup
                )
            except:
                bot.send_message(message.chat.id, f"⚠️ Не удалось уведомить пользователя {uid}")

            # Логирование
            admin_info = f"{message.from_user.first_name} {message.from_user.last_name or ''}"
            balance_log = (
                f"💰 ПОПОЛНЕНИЕ БАЛАНСА (команда)\n"
                f"👑 Админ: {admin_info}\n"
                f"👤 Пользователь: ID {uid}\n"
                f"💵 Сумма: {amt}₽\n"
                f"💳 Новый баланс: {new_balance}₽"
            )
            send_log_message(balance_log)

        else:
            bot.send_message(message.chat.id, "❌ Ошибка при пополнении")

    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Ошибка: {e}")




def check_old_payments():
    """Проверяет старые pending платежи и помечает их как expired"""
    try:
        conn = sqlite3.connect('balance_service/balance.db')
        cursor = conn.cursor()
        
        # Помечаем платежи старше 2 часов как expired
        cursor.execute('''
            UPDATE payments 
            SET status = 'expired' 
            WHERE status = 'pending' 
            AND datetime(created_at) < datetime('now', '-2 hours')
        ''')
        
        expired_count = cursor.rowcount
        if expired_count > 0:
            print(f"📋 Помечено {expired_count} просроченных платежей")
        
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"❌ Ошибка проверки старых платежей: {e}")

def schedule_payment_checker():
    """Запускает периодическую проверку платежей"""
    schedule.every(30).minutes.do(check_old_payments)
    
    while True:
        schedule.run_pending()
        time.sleep(60)

# Запускаем в отдельном потоке
payment_checker_thread = threading.Thread(target=schedule_payment_checker, daemon=True)
payment_checker_thread.start()

def sync_balance_with_site(user_id, amount):
    try:
        url = "https://fenix-osint.site/update_balance.php"
        data = {
            "user_id": user_id,
            "amount": amount
        }
        r = requests.post(url, data=data, timeout=10)
        print("📡 Синхронизация с сайтом:", r.text)
    except Exception as e:
        print("❌ Ошибка синхронизации с сайтом:", e)


def add_subscription(user_id, days):
    conn = sqlite3.connect('database.db')
    c = conn.cursor()

    c.execute("INSERT OR REPLACE INTO subscriptions (user_id, expires_at) VALUES (?, ?)", (user_id,))
    row = c.fetchone()

    now = datetime.now()
    if row and row[0]:  
        current_end = datetime.strptime(row[0], "%Y-%m-%d %H:%M:%S")
        if current_end > now:
            # если подписка активна – прибавляем
            new_end = current_end + timedelta(days=days)
        else:
            # если подписка закончилась – начинаем с сегодняшнего дня
            new_end = now + timedelta(days=days)
    else:
        # если подписки вообще не было
        new_end = now + timedelta(days=days)

    c.execute("UPDATE users SET subscription_end = ? WHERE user_id = ?", (new_end.strftime("%Y-%m-%d %H:%M:%S"), user_id))
    conn.commit()
    conn.close()

@bot.callback_query_handler(func=lambda call: call.data.startswith("hide_notification_"))
def hide_notification(call):
    try:
        # Удаляем сообщение с уведомлением
        bot.delete_message(call.message.chat.id, call.message.message_id)
        bot.answer_callback_query(call.id, "Уведомление скрыто")
    except Exception as e:
        print(f"Ошибка при скрытии уведомления: {e}")
        bot.answer_callback_query(call.id, "Не удалось скрыть уведомление")

def ensure_stats_table_exists():
    """Создает таблицу stats, если она не существует"""
    try:
        conn = sqlite3.connect('stats.db')
        c = conn.cursor()
        c.execute('''
            CREATE TABLE IF NOT EXISTS stats (
                base_name TEXT PRIMARY KEY,
                record_count INTEGER
            )
        ''')
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Ошибка создания таблицы stats: {e}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("hide_notification"))
def hide_notification(call):
    try:
        # Удаляем сообщение с уведомлением
        bot.delete_message(call.message.chat.id, call.message.message_id)
        bot.answer_callback_query(call.id, "✅ Уведомление скрыто")
    except Exception as e:
        print(f"Ошибка при скрытии уведомления: {e}")
        try:
            bot.answer_callback_query(call.id, "❌ Не удалось скрыть уведомление")
        except:
            pass
    
@bot.callback_query_handler(func=lambda call: call.data == "hide_msg")
def hide_message(call):
    try:
        bot.delete_message(call.message.chat.id, call.message.message_id)
    except:
        bot.answer_callback_query(call.id, "⚠️ Не удалось удалить сообщение")

SUBS_DB_PATH = os.path.join('subscriptions_service', 'subs.db')

def get_subs_conn():
    """Возвращает подключение к БД подписок с логом"""
    os.makedirs('subscriptions_service', exist_ok=True)
    conn = sqlite3.connect(SUBS_DB_PATH)
    return conn

def init_subscription_db():
    """Создаёт таблицу подписок, если её нет"""
    conn = get_subs_conn()
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS subscriptions (
            user_id INTEGER PRIMARY KEY,
            expires_at INTEGER NOT NULL
        )
    ''')
    conn.commit()
    conn.close()
    print(f"✅ Таблица подписок проверена/создана в {SUBS_DB_PATH}")

def is_subscribed(user_id: int) -> bool:
    """Проверка активна ли подписка"""
    try:
        init_subscription_db()
        conn = get_subs_conn()
        cursor = conn.cursor()
        cursor.execute("SELECT expires_at FROM subscriptions WHERE user_id = ?", (user_id,))
        result = cursor.fetchone()
        conn.close()
        if result:
            return time.time() < result[0]
        return False
    except Exception as e:
        print(f"[Ошибка проверки подписки]: {e}")
        return False

def add_subscription(user_id: int, duration_seconds: int) -> bool:
    """Добавление / продление подписки"""
    try:
        init_subscription_db()
        conn = get_subs_conn()
        cursor = conn.cursor()
        expires_at = int(time.time()) + duration_seconds
        cursor.execute('''
            INSERT OR REPLACE INTO subscriptions (user_id, expires_at)
            VALUES (?, ?)
        ''', (user_id, expires_at))
        conn.commit()
        conn.close()
        print(f"✅ Подписка для {user_id} до {datetime.fromtimestamp(expires_at)}")
        return True
    except Exception as e:
        print(f"[Ошибка добавления подписки]: {e}")
        return False

def send_log_message(text):
    """Отправляет сообщение в группу логов"""
    try:
        bot.send_message(LOG_GROUP_ID, text)
    except Exception as e:
        print(f"❌ Ошибка отправки лога: {e}")


if __name__ == "__main__":
    main_menu()
    init_databases()   